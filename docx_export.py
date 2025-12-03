"""
Модуль експорту звіту у формат DOCX (Microsoft Word).
Зберігає логіку візуалізації (діаграми) та структуру таблиць,
аналогічну PDF-звіту.
"""

import io
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import textwrap

from classification import QuestionInfo, QuestionType
from summary import QuestionSummary
from typing import List

# --- КОНСТАНТИ ДЛЯ ДІАГРАМ (Ті самі, що в PDF) ---
CHART_DPI = 150
FONT_SIZE_BASE = 11
BAR_WIDTH = 0.6

def create_chart_image(qs: QuestionSummary) -> io.BytesIO:
    """
    Генерує зображення діаграми (Matplotlib) і повертає його як байтовий потік.
    """
    plt.clf() # Очищення попереднього малюнка
    plt.rcParams.update({'font.size': FONT_SIZE_BASE})
    
    labels = qs.table["Варіант відповіді"].astype(str).tolist()
    values = qs.table["Кількість"]
    wrapped_labels = [textwrap.fill(l, 40) for l in labels]

    # Вибір типу та розміру діаграми (Логіка PDF)
    if qs.question.qtype == QuestionType.SCALE:
        # Стовпчикова: компактна по висоті
        fig = plt.figure(figsize=(10, 4.0))
        bars = plt.bar(wrapped_labels, values, color='#4F81BD', width=BAR_WIDTH)
        plt.ylabel('Кількість')
        plt.grid(axis='y', linestyle='--', alpha=0.5)
        plt.xticks(rotation=0)
        
        for bar in bars:
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                     f'{int(height)}', ha='center', va='bottom', fontweight='bold')
    else:
        # Кругова: висока для легенди
        fig = plt.figure(figsize=(10, 6.5))
        colors = ['#4F81BD', '#C0504D', '#9BBB59', '#8064A2', '#4BACC6', '#F79646']
        c_arg = colors[:len(values)] if len(values) <= len(colors) else None
        
        wedges, texts, autotexts = plt.pie(
            values, labels=None, autopct='%1.1f%%', startangle=90,
            pctdistance=0.8, colors=c_arg, radius=1.2,
            textprops={'fontsize': FONT_SIZE_BASE}
        )
        
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_weight('bold')
            import matplotlib.patheffects as path_effects
            autotext.set_path_effects([path_effects.withStroke(linewidth=2, foreground='#333333')])

        plt.axis('equal')
        
        cols = 2 if len(labels) > 3 else 1
        plt.legend(
            wrapped_labels,
            loc="upper center",
            bbox_to_anchor=(0.5, -0.05),
            ncol=cols,
            frameon=False
        )

    plt.tight_layout()
    
    # Зберігаємо в буфер пам'яті
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=CHART_DPI, bbox_inches='tight')
    plt.close(fig)
    img_stream.seek(0)
    return img_stream

def build_docx_report(
    original_df: pd.DataFrame,
    sliced_df: pd.DataFrame,
    summaries: List[QuestionSummary],
    range_info: str,
) -> bytes:
    """
    Створює DOCX файл і повертає його байти.
    """
    doc = Document()
    
    # --- 1. Налаштування стилів ---
    # Базовий шрифт
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # --- 2. Титульна сторінка ---
    title = doc.add_heading('Звіт про результати опитування', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Всього анкет у файлі: {len(original_df)}")
    doc.add_paragraph(f"Оброблено анкет: {len(sliced_df)}")
    doc.add_paragraph(f"Діапазон обробки: {range_info}")
    doc.add_paragraph("_" * 50) # Розділювач
    doc.add_paragraph("") # Відступ

    # --- 3. Основний цикл по питаннях ---
    for qs in summaries:
        # Заголовок питання (виділений фоном в Word складно зробити стилем, тому просто Жирний Heading)
        q_title = f"{qs.question.code}. {qs.question.text}"
        heading = doc.add_heading(q_title, level=2)
        
        # Перевірка на дані
        if qs.table.empty:
            doc.add_paragraph("Немає даних для відображення або текстові відповіді.")
            doc.add_paragraph("")
            continue

        # --- Таблиця ---
        # add_table(rows, cols)
        # rows = к-сть даних + 1 (хедер)
        table = doc.add_table(rows=len(qs.table)+1, cols=3)
        table.style = 'Table Grid' # Стиль з рамками
        table.autofit = True
        
        # Хедер
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Варіант відповіді"
        hdr_cells[1].text = "Кількість"
        hdr_cells[2].text = "%"
        
        # Жирний шрифт для хедера
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Заповнення даними
        for i, row in enumerate(qs.table.itertuples(index=False)):
            row_cells = table.rows[i+1].cells
            
            # Варіант
            row_cells[0].text = str(row[0])
            
            # Кількість
            row_cells[1].text = str(row[1])
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Відсоток
            row_cells[2].text = str(row[2])
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("") # Відступ після таблиці

        # --- Діаграма ---
        try:
            image_stream = create_chart_image(qs)
            # Вставляємо картинку. Width=Inches(6) — це ширина сторінки А4 з полями
            doc.add_picture(image_stream, width=Inches(6.0))
            
            # Центрування останнього параграфа (картинки)
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph("") # Відступ після діаграми
        except Exception as e:
            doc.add_paragraph(f"[Помилка побудови діаграми: {e}]")

        # Розрив сторінки перед наступним питанням (опціонально, але для звітності гарно)
        # Щоб не було "каші", додаємо розрив, якщо це не останнє питання
        if qs != summaries[-1]:
             # Можна додати doc.add_page_break(), якщо хочете кожне питання з нової сторінки
             # Або просто лінію
             doc.add_paragraph("_" * 20)
             doc.add_paragraph("")

    # --- 4. Збереження в буфер ---
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()